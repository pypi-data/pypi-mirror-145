"""Documentation about pynterlinear"""
import logging
import os
import re
from docx import Document
from docx import shared


logging.getLogger(__name__).addHandler(logging.NullHandler())

__author__ = "Florian Matter"
__email__ = "florianmatter@gmail.com"
__version__ = "0.0.2"


# These symbols could potentially be used to split up morphemes.
# Some of them are standard, some not.
delimiters = [
    "-",
    "–",
    ".",
    "=",
    ":",
    "*",
    "~",
    "<",
    ">",
    "[",
    "]",
    "(",
    ")",
    "/",
    "\\",
]
morpheme_delimiters = ["-", "=", "~"]

# Create a hash of common glossing abbreviations and their meaning.
glossing_abbrevs = {}
data_dir = os.path.join(os.path.dirname(__file__), "data")
fn = os.path.join(data_dir, "glossing.txt")
with open(fn, "r", encoding="utf-8") as f:
    raw_glosses = f.read()
for entry in raw_glosses.split("\n"):
    glossing_abbrevs[entry.split("\t")[0]] = entry.split("\t")[1]

# This is used for keeping track of abbreviations that came up in the analysis
# but which are not in glossing_abbrevs.
# This script is primarily geared towards the LaTeX package glossingtool.sty,
# which has the same predefined glossing abbreviations.
unknown_abbrevs = []


def get_unknown_abbrevs():
    return unknown_abbrevs


def get_delims():
    return delimiters


def get_all_abbrevs():
    return glossing_abbrevs


def pad_ex(obj, gloss, as_tuple=False, as_list=False):
    out_obj = []
    out_gloss = []
    if not as_list:
        zipp = zip(obj.split(" "), gloss.split(" "))
    else:
        zipp = zip(obj, gloss)
    for o, g in zipp:
        diff = len(o) - len(g)
        if diff < 0:
            o += " " * -diff
        else:
            g += " " * diff
        out_obj.append(o)
        out_gloss.append(g)
    if as_tuple:
        return "  ".join(out_obj).strip(" "), "  ".join(out_gloss).strip(" ")
    return "  ".join(out_obj).strip(" ") + "\n" + "  ".join(out_gloss).strip(" ")


# Can be used to quickly look up the proper abbreviation for something.
def search_abbrev(term):
    output = []
    for key, abbrev in glossing_abbrevs.items():
        if term in abbrev:
            output.append(f"{key}\t{abbrev}")
    return "\n".join(output)


# This takes an object word and a gloss and returns a form--meaning hash
def get_morphemes_from_word(object_word, gloss_word):
    for delim in morpheme_delimiters:
        object_word = object_word.replace(delim, " ")
    for delim in morpheme_delimiters:
        gloss_word = gloss_word.replace(delim, " ")
    forms = object_word.split(" ")
    glosses = gloss_word.split(" ")
    if len(forms) != len(glosses):
        print(f"{object_word}: mismatch in morpheme number!")
    return dict(zip(forms, glosses))


# This takes a whole example (hash) and returns a form--meaning hash
def get_morphemes(**example):
    pairings = {}
    l1, l2 = len(example["obj"]), len(example["gloss"])
    if l1 != l2:
        print(f"""{example["id"]}: mismatch in word number! {l1} vs {l2} words""")
    for obj, gloss in zip(example["obj"], example["gloss"]):
        new_entries = get_morphemes_from_word(obj, gloss)
        pairings = {**pairings, **new_entries}
    return pairings


# Splits a word up into morphemes and delimiters
def split_word(word):
    output = []
    char_list = list(word)
    for char in char_list:
        if len(output) == 0 or (char in delimiters or output[-1] in delimiters):
            output.append(char)
        else:
            output[-1] += char
    return output


# Takes an uppercase string like 1SG and breaks it up into
# known abbreviations like ["1", "SG"]
def get_glossing_combination(input_string):
    output = []
    temp_text = ""
    for i, char in enumerate(list(input_string)):
        if re.match(r"[1-3]+", char):
            if i < len(input_string) - 1 and input_string[i + 1] == "+":
                temp_text += char
            elif input_string[i - 1] == "+":
                temp_text += char
                output.append(temp_text)
                temp_text = ""
            else:
                if temp_text != "":
                    output.append(temp_text)
                output.append(char)
                temp_text = ""
        else:
            temp_text += char
    if temp_text != "":
        output.append(temp_text)
    for morpheme in output:
        if (
            morpheme.lower() not in glossing_abbrevs
            and morpheme.lower() not in unknown_abbrevs
        ):
            unknown_abbrevs.append(morpheme.lower())
    return output


def is_glossing_candidate(part, parts, j):
    return (
        part == part.upper()
        and part not in delimiters
        and part != "?"
        and not (
            len(part) == 1
            and not re.match(r"\d", part)  # is it only one capital letter?
            and (
                # len(parts) == j+1 #are we at the end of the word?
                # or (
                # or are there more characters?
                len(parts) == j + 2
                and parts[j + 1] in ["."]  # and is the next character a period?
            )
        )
    )


# Creates expex code with \gl{} for glossing abbreviations
def get_expex_code(input_string):
    input_string = input_string.replace("\\", "\\textbackslash()")
    words_list = input_string.split(" ")
    # iterate through words of glossing line
    for i, word in enumerate(words_list):
        # Transform _X_ to subscript beforehand
        word = re.sub(r"_([^AP_]+?)_", r"\\\\textsubscript{\g<1>}", word)
        output = " "
        # take proper nouns into account
        if len(word) == 2 and word[0] == word[0].upper() and word[1] == ".":
            output += word
        else:
            parts = split_word(word)
            for j, part in enumerate(parts):
                if is_glossing_candidate(part, parts, j):
                    if part.lower() in glossing_abbrevs:
                        output += f"\\gl{{{part.lower()}}}"
                    # take care of numbered genders
                    elif part[0] == "G" and re.match(r"\d", part[1:]):
                        output += f"\\gl{{g}}{part[1:]}"
                    else:
                        for extracted_morpheme in get_glossing_combination(part):
                            output += f"\\gl{{{extracted_morpheme.lower()}}}"
                else:
                    output += part
        words_list[i] = output[1:]
    gloss_text_upcased = " ".join(words_list)
    gloss_text_upcased = (
        gloss_text_upcased.replace("~", "\\glosstilde{}")
        .replace("_", "_")
        .replace("_a_", "_a_")
        .replace("_p_", "_p_")
        .replace("textbackslash()", "textbackslash{}")
    )
    return gloss_text_upcased


# This function takes a list of dictionaries and produces ExPex code
# Use for_beamer if you want to insert a \glottolink command
# For this, a key glottocode is needed in the example hash.
def convert_to_expex(
    examples,
    for_beamer=False,
    from_corpus=False,
    pextag=None,
    latex_labels=False,
    multicols=False,
    no_surf=False,
):
    # See what languages and sources we're dealing with and whether it makes
    # sense to just print them once
    languages = []
    sources = []
    lengths = []
    for example in examples:
        if "language" in example.keys():
            languages.append(example["language"])
        if "source" in example.keys():
            sources.append(example["source"].split("[")[0])
        lengths.append(len(example["obj"]) + 5)
        lengths.append(len(example["gloss"]) + 5)
        lengths.append(len(example["trans"]))
    # If there is only one language, we only need to print it once
    same_language = len(set(languages)) < 2
    # If there is only one source, we only need to print it once
    same_source = len(set(sources)) == 1
    if pextag:
        pex = True
        output = f"\\pex<{pextag}>"
    else:
        pex = False
        output = ""
    len_threshold = max(lengths)
    multicols = (
        pex and multicols and len_threshold < 35 and same_language and same_source
    )
    # If there is one and only language specified, we print it right at the
    # start, after \(p)ex
    language_string = ""
    if same_language and "language" in examples[0].keys():
        # Print \glottolink if for_beamer is set to True
        if for_beamer:
            gl, lg = examples[0]["glottocode"], examples[0]["language"]
            language_string = f"""\\glottolink{{{gl}}}{{{lg}}}"""
        elif same_language and from_corpus:
            language_string = ""
        else:
            language_string = examples[0]["language"]
        # Add the language name to the output
        output += language_string
    else:
        same_source = False
    if "parnote" in examples[0].keys():
        parnote = examples[0]["parnote"]
    else:
        parnote = ""
    # Same logic as before, if there is only one source (for the one
    # language), we add it only once, at the start
    if same_source:
        logging.debug(f"All subexamples of pex {pextag} have the same source.")
        logging.debug(examples)
        # Gather the page (ranges) given in the single sources in one list
        page_string = []
        for example in examples:
            page_string.append(example["source"].split("[")[1].split("]")[0])
        if len(list(set(page_string))) == 1:
            page_string = page_string[0]
        else:
            page_string = ", ".join(page_string)
        # Print the list joined by commas
        if sources[0] != "pc":
            output += f" \\parencite[{parnote}][{page_string}]{{{sources[0]}}}"
        else:
            output += " \\perscommpar{{{page_string}}}"
    if latex_labels and pex:
        output += f"\\exl{{{pextag}}}"
    # beamer for some reason wants a line break after \pex, but absolutely
    # not after \ex; text documents want the opposite
    if (for_beamer and pex) or (not for_beamer and not pex and language_string != ""):
        output += "\\\\"
    # If we're dealing with subexamples, we need a linebreak before the \a
    if pex:
        output += "\n"
        # Do we want multicols?
        if multicols:
            output += "\\begin{multicols}{2}\n"
    # part_text holds everything between \a (if present) and \endgl.  It is used
    # even in the case of a single example
    for example in examples:
        # We simply add nothing if we're dealing with a single example,
        # otherwise we add an \a with a tag
        if pex:
            part_text = f"""\\a<{example["id"]}>"""
        else:
            part_text = f"""\\ex<{example["id"]}>"""
        # If there is a source, we extract the details
        if "source" in example.keys():
            pages = example["source"].split("[")[1].split("]")[0]
            source_key = example["source"].split("[")[0]
            if "parnote" in example.keys():
                parnote = example["parnote"]
            else:
                parnote = ""
        if from_corpus:
            speaker = ""
            start = ""
            end = ""
            part = ""
            text_id = ""
            if "speaker" in example.keys():
                speaker = example["speaker"]
            if "start" in example.keys():
                start = example["start"]
            if "end" in example.keys():
                end = example["end"]
            if "part" in example.keys():
                part = example["part"]
            if "text_id" in example.keys():
                text_id = example["text_id"]
        # If we're dealing with different languages, we want to print every one,
        # after the respective \a
        if not same_language:
            # Again with the glottolink string
            if for_beamer:
                gl, lg = examples[0]["glottocode"], examples[0]["language"]
                language_string = f"""\\glottolink{{{gl}}}{{{lg}}}"""
            else:
                language_string = example["language"]
            part_text += " " + language_string
            # And if we're not dealing with a single source, we add all the
            # sources seperately
            if not same_source and not from_corpus:
                if source_key != "pc":
                    part_text += f" \\parencite[{parnote}][{pages}]{{{source_key}}}"
                else:
                    part_text += f" \\perscommpar{{{pages}}}"
        # Text documents also want a linebreak after \a with text behind it,
        # beamer doesn't
        if not for_beamer and not same_language:
            print("HELLO")
            part_text += "\\\\"
        # Line break before \begingl
        part_text += "\n"
        # If there is a surface form, we add it, otherwise it's just an empty
        # string
        if (
            "surface" not in example or example["surface"] == example["obj"] or no_surf
        ):  # or "-" in example["surface"]:
            surface_string = ""
        else:
            if "[" in example["surface"] and "]" in example["surface"]:
                surface_string = (
                    f"""\\glpreamble \\normalfont {example["surface"]}//\n"""
                )
            else:
                surface_string = f"""\\glpreamble {example["surface"]}//\n"""
        # allow custom glc contents
        glcstring = ""
        for k in example.keys():
            if "glc" in k:
                glcstring += "\n\\glc " + example[k] + "//"
        # Get the \gl{}-ified interlinear gloss
        gloss_text_markup = get_expex_code(example["gloss"])
        # Put together the interlinear text
        mod_obj = example["obj"].replace("~", "\\glosstilde{}")
        part_text += f"""\\begingl
{surface_string}
\\gla {mod_obj}//
\\glb {gloss_text_markup}//{glcstring}
\\glft \\qu{{{example["trans"]}}}"""
        # If we're dealing with one language (which we printed after \pex), but
        # with different sources, we wanna print the sources after the
        # translation
        if same_language and not same_source and "source" in example.keys():
            if source_key != "pc":
                part_text += f" \\parencite[{parnote}][{pages}]{{{source_key}}}"
            else:
                part_text += f" \\perscommpar{{{pages}}}"
        elif from_corpus:
            part_text += f" \\textref[speaker={speaker}, start={start}, \
             end={end}, part={part}]{{{text_id}}}"
        # Finish interlinear text
        part_text += "//"
        if latex_labels and pex:
            part_text += f"""\\exl{{{pextag}.{example["id"]}}}"""
        elif latex_labels and not pex:
            part_text += f"""\\exl{{{example["id"]}}}"""
        part_text += "\n\\endgl"
        if not pex and len(examples) > 1:
            part_text += "\n\\xe"
        # Add to output
        output += part_text + "\n"
    if multicols:
        output += "\\end{multicols}\n"
    # Finish \ex…\xe
    if pex:
        output += "\\xe"
    return output


def convert_to_word(examples, use_tables=True, filename="csv2word_export.docx"):
    def get_running_number_tables(document):
        for i in range(1, len(document.tables) + 1):
            topright = document.tables[-i].rows[0].cells[0].text
            search = re.search("\\((.*)\\)", topright)
            if search:
                return int(search.group(1))
        return 0

    def get_running_number_tabs(document):
        for i in range(1, len(document.paragraphs) + 1):
            partext = document.paragraphs[-i].text
            search = re.search("\\((.*)\\)", partext)
            if search:
                return int(search.group(1))
        return 0

    if filename in os.listdir("."):
        document = Document(filename)
        if use_tables:
            running_number = get_running_number_tables(document) + 1
        else:
            running_number = get_running_number_tabs(document) + 1
    else:
        document = Document()
        running_number = 1

    if use_tables:
        xs = 1
        if len(examples) > 1:
            xs += 1
        exhe = 3
        for exno, example in enumerate(examples):
            obj_words = example["obj"].split(" ")
            gloss_words = example["gloss"].split(" ")
            trans = example["trans"]
            table = document.add_table(rows=exhe, cols=len(obj_words) + xs)
            # Only for development purposes
            # table.style = "Table Grid"
            for i, obj_word in enumerate(obj_words):
                table.rows[0].cells[i + xs].paragraphs[0].add_run(
                    obj_word
                ).italic = True
            for i, gloss_word in enumerate(gloss_words):
                if (
                    len(gloss_word) == 2
                    and gloss_word[0] == gloss_word[0].upper()
                    and gloss_word[1] == "."
                ):
                    tt = (
                        table.rows[1]
                        .cells[i + xs]
                        .paragraphs[0]
                        .add_run(f"{gloss_word}")
                    )
                else:
                    morphemes = split_word(gloss_word)
                    for morpheme in morphemes:
                        # take proper nouns into account
                        if (
                            morpheme == morpheme.upper()
                            and morpheme not in delimiters
                            and morpheme != "?"
                        ):
                            tt = (
                                table.rows[1]
                                .cells[i + xs]
                                .paragraphs[0]
                                .add_run(f"{morpheme.lower()}")
                            )
                            tt.font.small_caps = True
                        else:
                            tt = (
                                table.rows[1]
                                .cells[i + xs]
                                .paragraphs[0]
                                .add_run(f"{morpheme}")
                            )
            for column in table.columns:
                for cell in column.cells:
                    cell._tc.tcPr.tcW.type = "auto"
            transcell = table.rows[-1].cells[0 + xs].merge(table.rows[-1].cells[-1])
            transcell.text = f"‘{trans}’"
            table.rows[0].cells[0].allow_autofit = False
            table.rows[0].cells[0].width = shared.Cm(1)
            if exno == 0:
                table.rows[0].cells[0].text = f"({running_number})"
            if len(examples) > 1:
                table.rows[0].cells[1].text = f"{chr(97 + exno)}."
            # Remove random 10pt spacing in EVERY CELL
            for row in table.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        par.paragraph_format.space_after = shared.Pt(0)
            par = document.add_paragraph("")
            par.paragraph_format.space_after = shared.Pt(0)
    else:
        gloss_para = document.add_paragraph(f"({running_number})")
        for exno, example in enumerate(examples):
            if len(examples) > 1:
                sub_ex_label = chr(97 + exno) + ".\t"
                tabstops = [1, 2]
            else:
                sub_ex_label = ""
                tabstops = [1.25]
            gloss_para.add_run(f"\t{sub_ex_label}")
            mod_obj = example["obj"].replace(" ", "\t")
            obj_run = gloss_para.add_run(f"""{mod_obj}\n\t\t""")
            obj_run.italic = True
            for i, gloss_word in enumerate(example["gloss"].split(" ")):
                if (
                    len(gloss_word) == 2
                    and gloss_word[0] == gloss_word[0].upper()
                    and gloss_word[1] == "."
                ):
                    add_text = gloss_para.add_run(gloss_word)
                else:
                    morphemes = split_word(gloss_word)
                    for morpheme in morphemes:
                        # take proper nouns into account
                        if (
                            morpheme == morpheme.upper()
                            and morpheme not in delimiters
                            and morpheme != "?"
                        ):
                            add_text = gloss_para.add_run(morpheme.lower())
                            add_text.font.small_caps = True
                        else:
                            add_text = gloss_para.add_run(morpheme)
                if i < len(example["gloss"].split(" ")) - 1:
                    gloss_para.add_run("\t")
            gloss_para.add_run(
                f"""\n\t\t‘{example["trans"]}’"""
            )  # , sub_example["source"]))
            tab_stops = gloss_para.paragraph_format.tab_stops
            for tabstop in tabstops:
                tab_stops.add_tab_stop(shared.Cm(tabstop))
            gloss_para = document.add_paragraph()
    document.save(filename)
